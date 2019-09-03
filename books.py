from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QWidget, QMessageBox
import sys, re
from PyQt5.QtCore import QDateTime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH



class Ui_Dialog(object):

    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.resize(800, 800)

        # line edit
        self.lineEdit_title = QtWidgets.QLineEdit(Dialog)
        self.lineEdit_title.setGeometry(QtCore.QRect(90, 20, 300, 20))
        self.lineEdit_title.setObjectName("lineEdit_title")
        self.lineEdit_author = QtWidgets.QLineEdit(Dialog)
        self.lineEdit_author.setGeometry(QtCore.QRect(90, 50, 150, 20))
        self.lineEdit_author.setObjectName("lineEdit_author")
        self.lineEdit_trans = QtWidgets.QLineEdit(Dialog)
        self.lineEdit_trans.setGeometry(QtCore.QRect(330, 50, 150, 20))
        self.lineEdit_trans.setObjectName("lineEdit_trans")
        self.label = QtWidgets.QLabel(Dialog)
        self.label.setGeometry(QtCore.QRect(20, 20, 60, 20))
        self.lineEdit_company = QtWidgets.QLineEdit(Dialog)
        self.lineEdit_company.setGeometry(QtCore.QRect(570, 50, 120, 20))
        self.lineEdit_company.setObjectName("lineEdit_company")

        # label
        font = QtGui.QFont()
        font.setPointSize(11)
        self.label.setFont(font)
        self.label.setFocusPolicy(QtCore.Qt.NoFocus)
        self.label.setAlignment(QtCore.Qt.AlignCenter)
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(Dialog)
        self.label_2.setGeometry(QtCore.QRect(20, 50, 60, 20))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.label_2.setFont(font)
        self.label_2.setAlignment(QtCore.Qt.AlignCenter)
        self.label_2.setObjectName("label_2")
        self.label_3 = QtWidgets.QLabel(Dialog)
        self.label_3.setGeometry(QtCore.QRect(270, 50, 60, 20))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.label_3.setFont(font)
        self.label_3.setAlignment(QtCore.Qt.AlignCenter)
        self.label_3.setObjectName("label_3")
        self.label_4 = QtWidgets.QLabel(Dialog)
        self.label_4.setGeometry(QtCore.QRect(510, 50, 60, 20))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.label_4.setFont(font)
        self.label_4.setAlignment(QtCore.Qt.AlignCenter)
        self.label_4.setObjectName("label_4")
        self.label_5 = QtWidgets.QLabel(Dialog)
        self.label_5.setGeometry(QtCore.QRect(470, 20, 30, 30))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.label_5.setFont(font)
        self.label_5.setAlignment(QtCore.Qt.AlignCenter)
        self.label_5.setObjectName("label_5")
        self.label_6 = QtWidgets.QLabel(Dialog)
        self.label_6.setGeometry(QtCore.QRect(560, 20, 60, 30))
        self.label_6.setObjectName("label_6")


        # spin box
        self.spinBox = QtWidgets.QSpinBox(Dialog)
        self.spinBox.setGeometry(QtCore.QRect(505, 25, 40, 20))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.spinBox.setFont(font)
        self.spinBox.setMinimum(1)
        self.spinBox.setMaximum(5)
        self.spinBox.setProperty("value", 3)
        self.spinBox.setObjectName("spinBox")
        self.spinBox.valueChanged.connect(self.clickstar)

        # etc
        self.textEdit = QtWidgets.QTextEdit(Dialog)
        self.textEdit.setGeometry(QtCore.QRect(10, 90, 781, 650))
        self.textEdit.setObjectName("textEdit")
        self.pushButton = QtWidgets.QPushButton(Dialog)
        self.pushButton.setGeometry(QtCore.QRect(10, 760, 781, 31))
        font = QtGui.QFont()
        font.setPointSize(13)
        self.pushButton.setFont(font)
        self.pushButton.setObjectName("pushButton")
        self.pushButton.clicked.connect(self.clicksave)

        self.retranslateUi(Dialog)
        QtCore.QMetaObject.connectSlotsByName(Dialog)


    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "독후감"))
        self.label.setText(_translate("Dialog", "제목"))
        self.label_2.setText(_translate("Dialog", "지은이"))
        self.label_3.setText(_translate("Dialog", "옮긴이"))
        self.label_4.setText(_translate("Dialog", "출판사"))
        self.label_5.setText(_translate("Dialog", "별"))
        self.label_6.setText(_translate("Dialog", "* * *"))
        self.pushButton.setText(_translate("Dialog", "저장"))


    def clickstar(self):
        num = self.spinBox.value()
        star = "* "*num
        self.label_6.setText(star)


    def clicksave(self):

        dat = QDateTime.currentDateTime().toString('yyyyMMdd')
        date = dat[:4]+'년 '+dat[4:6]+'월 '+dat[6:]+'일'

        title = self.lineEdit_title.text()
        star = "* "* self.spinBox.value()
        author = self.lineEdit_author.text()
        trans = self.lineEdit_trans.text()
        company = self.lineEdit_company.text()
        content = self.textEdit.toPlainText()


        # 빼먹으면 안되
        if len(title)==0:
            self.label.setText("제목써라")
        elif len(author)==0:
            self.label_2.setText("지은이써라")
        elif len(company)==0:
            self.label_4.setText("출판사써라")
        elif len(content)==0:
            self.textEdit.append("내용을 써라")
        else:
            document = Document()
            document.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph()

         #옮긴이가 없으면 뺌
            if len(trans) == 0 :
                document.add_paragraph('지은이 : '+author+'      '+'출판사 : '+company).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else :
                document.add_paragraph('지은이 : '+author+'      '+'옮긴이 : '+trans+'      '+'출판사 : '+company).alignment = WD_ALIGN_PARAGRAPH.RIGHT

            document.add_paragraph('별점 : '+star+'          '+str(date)).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            document.add_paragraph()
            document.add_paragraph(content)
            document.save('D:/books/'+dat+title+'.docx')

            self.lineEdit_title.clear()
            self.spinBox.setProperty("value", 3)
            self.lineEdit_author.clear()
            self.lineEdit_trans.clear()
            self.lineEdit_company.clear()
            self.textEdit.clear()
            self.textEdit.append(title+" 저장했어")


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    Dialog = QtWidgets.QDialog()
    ui = Ui_Dialog()
    ui.setupUi(Dialog)
    Dialog.show()
    sys.exit(app.exec_())
